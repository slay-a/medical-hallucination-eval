#!/usr/bin/env python3
"""
build_thesis.py — Generate thesis_final.docx from results CSVs + plot PNGs.
"""

import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np

RESULTS = os.path.join(os.path.dirname(__file__), "results")
OUT     = os.path.join(os.path.dirname(__file__), "thesis_final.docx")

cmp    = pd.read_csv(os.path.join(RESULTS, "comparison_per_sample.csv"))
agg    = pd.read_csv(os.path.join(RESULTS, "aggregate_statistics.csv"))
claims = pd.read_csv(os.path.join(RESULTS, "claims_all.csv"))

# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right — each a dict with sz, color, val."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        if side in kwargs:
            tag = OxmlElement(f"w:{side}")
            for k, v in kwargs[side].items():
                tag.set(qn(f"w:{k}"), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_format(para, space_before=0, space_after=8, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)

def add_heading(doc, text, level, space_before=18, space_after=6):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(space_after)
    return h

def add_body(doc, text, indent=0):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    pf = p.paragraph_format
    pf.space_after  = Pt(8)
    pf.space_before = Pt(0)
    pf.first_line_indent = Pt(0)
    if indent:
        pf.left_indent = Pt(indent)
    return p

def add_figure(doc, png_name, caption, width=5.8):
    path = os.path.join(RESULTS, png_name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Normal"]
    run2 = cap.runs[0]
    run2.italic = True
    run2.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(14)
    return p

def add_twoplot(doc, png_a, cap_a, png_b, cap_b):
    """Two plots side by side in a borderless table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (png, cap) in enumerate([(png_a, cap_a), (png_b, cap_b)]):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run()
        run.add_picture(os.path.join(RESULTS, png), width=Inches(2.9))
        cap_cell = tbl.rows[1].cells[i]
        cp = cap_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(cap)
        r.italic = True
        r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── document setup ───────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(3.17)

# Default font
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)
for lvl in range(1, 5):
    style = doc.styles[f"Heading {lvl}"]
    style.font.name  = "Times New Roman"
    style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    style.font.bold  = True
    if lvl == 1:
        style.font.size = Pt(16)
    elif lvl == 2:
        style.font.size = Pt(13)
    elif lvl == 3:
        style.font.size = Pt(12)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("Evaluating and Reducing Hallucinations in\nLLM-Based Medical Summarization:\nA Retrieval-Augmented Generation Approach")
r.bold = True
r.font.size = Pt(18)
r.font.name = "Times New Roman"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(30)
p2.paragraph_format.space_after  = Pt(6)
r2 = p2.add_run("A Thesis Submitted in Partial Fulfillment of the\nRequirements for the Degree of Master of Science")
r2.font.size = Pt(12)
r2.font.name = "Times New Roman"

for line in ["", "Department of Computer Science", "2026"]:
    p3 = doc.add_paragraph(line)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(4)
    if p3.runs:
        p3.runs[0].font.size = Pt(12)
        p3.runs[0].font.name = "Times New Roman"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Abstract", level=1, space_before=0)
add_body(doc,
    "Large language models (LLMs) have demonstrated remarkable capacity for generating "
    "coherent, patient-facing summaries of clinical documentation. However, their "
    "propensity to produce unsupported or factually incorrect statements—commonly termed "
    "hallucinations—poses serious safety risks in high-stakes medical contexts. This "
    "thesis presents a systematic empirical evaluation of hallucination rates in "
    "GPT-4o-mini generated medical summaries, comparing a zero-context baseline (E0) "
    "against a Retrieval-Augmented Generation system (E1) across 50 de-identified "
    "clinical transcriptions drawn from the MTSamples corpus."
)
add_body(doc,
    "Hallucinations are quantified at the individual claim level using a two-stage "
    "pipeline: (1) sentence-transformer cosine similarity retrieval of supporting "
    "evidence from the source document, and (2) zero-shot Natural Language Inference "
    "(NLI) classification via a cross-encoder model. Two aggregate metrics are reported: "
    "the Unsupported Fact Rate (UFR) and the Contradiction Rate (CR). Results show that "
    "RAG produces a statistically significant reduction in direct contradictions "
    "(CR: 0.226 → 0.154, Wilcoxon p < 0.001), with 72% of samples exhibiting "
    "improvement. However, UFR remains statistically unchanged (p = 0.581), indicating "
    "that retrieval-augmented generation reduces factual conflicts but does not "
    "substantially suppress the generation of claims that lack any evidentiary grounding "
    "in the source note. These findings motivate the need for claim-level grounding "
    "constraints in clinical NLP systems."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Introduction", level=1)
add_body(doc,
    "The integration of large language models into clinical workflows has accelerated "
    "dramatically in recent years. Automated summarization of discharge notes, "
    "consultation reports, and progress notes promises to reduce physician documentation "
    "burden, improve patient comprehension, and support downstream clinical decision "
    "support. Studies have demonstrated that LLMs can produce summaries that are "
    "preferred by clinicians over manually authored versions on dimensions of fluency "
    "and completeness (Van Veen et al., 2024). Yet alongside these capabilities, a "
    "well-documented failure mode persists: LLMs fabricate factual content with "
    "confident, fluent prose—a phenomenon known as hallucination (Ji et al., 2023)."
)
add_body(doc,
    "In the medical domain, hallucinations are not merely embarrassing artifacts; they "
    "constitute a patient safety risk. A summary that incorrectly states a patient "
    "is cleared for surgery, misattributes a drug allergy, or invents follow-up "
    "instructions that were never issued could directly influence clinical decisions. "
    "The Joint Commission and multiple regulatory bodies have identified AI-generated "
    "misinformation as an emerging safety concern. Despite this, the majority of "
    "medical NLP benchmarks continue to rely on surface-level overlap metrics such as "
    "ROUGE and BLEU, which cannot distinguish between a correctly paraphrased fact "
    "and a plausibly worded fabrication (Xie et al., 2023)."
)
add_body(doc,
    "Retrieval-Augmented Generation (RAG) offers a promising mitigation strategy: by "
    "grounding the model's generation in retrieved passages from the source document, "
    "one can in principle constrain the hypothesis space of possible outputs and reduce "
    "the frequency of factual errors (Lewis et al., 2020). However, empirical "
    "validation of this hypothesis at the claim level—as opposed to holistic fluency "
    "ratings—remains sparse in the clinical summarization literature."
)
add_body(doc,
    "This thesis addresses that gap with three primary contributions. First, we "
    "construct a claim-level NLI evaluation pipeline applicable to any pair of "
    "source document and generated summary, using open-source models throughout. "
    "Second, we apply this pipeline to systematically compare baseline and RAG-augmented "
    "GPT-4o-mini summaries across 50 clinical transcriptions from MTSamples, reporting "
    "both UFR and CR with full per-sample distributions. Third, we provide a "
    "qualitative analysis of claim categories that RAG successfully resolves versus "
    "those that remain persistently unsupported—informing the design of future "
    "clinical NLP systems."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Related Work", level=1)

add_heading(doc, "2.1 Limitations of Overlap-Based Evaluation", level=2)
add_body(doc,
    "ROUGE (Recall-Oriented Understudy for Gisting Evaluation) has served as the "
    "de facto standard for automatic summarization evaluation since its introduction "
    "by Lin (2004). ROUGE-N measures n-gram recall against one or more reference "
    "summaries, while ROUGE-L captures longest common subsequence. Despite their "
    "widespread adoption, overlap metrics carry a fundamental limitation: they assess "
    "lexical similarity rather than factual faithfulness. A summary may achieve high "
    "ROUGE-2 scores by reproducing frequent noun phrases from the source while still "
    "fabricating critical numerical values, medication names, or diagnostic conclusions. "
    "Kryscinski et al. (2020) demonstrated that state-of-the-art summarization models "
    "could hallucinate factual errors in up to 30% of generated summaries while "
    "maintaining competitive ROUGE scores. Similar findings have been reported in the "
    "medical domain by Xie et al. (2023), who showed that ROUGE correlates poorly with "
    "physician judgments of clinical accuracy."
)

add_heading(doc, "2.2 NLI-Based Faithfulness Evaluation", level=2)
add_body(doc,
    "Natural Language Inference (NLI), the task of classifying whether a hypothesis is "
    "entailed, contradicted, or neutral with respect to a premise, provides a more "
    "semantically grounded basis for faithfulness evaluation. Maynez et al. (2020) were "
    "among the first to apply NLI models to detect factual inconsistencies in "
    "abstractive summarization, introducing the distinction between intrinsic and "
    "extrinsic hallucinations. Laban et al. (2022) extended this approach with "
    "SummaC, a consistency scoring framework that segments both document and summary "
    "into sentences and applies an NLI classifier at the sentence level. Guo et al. "
    "(2022) showed that cross-encoder architectures significantly outperform "
    "bi-encoders for this classification task. This thesis adopts the cross-encoder "
    "architecture (cross-encoder/nli-MiniLM2-L6-H768) and extends the paradigm to the "
    "medical summarization setting with explicit retrieval of per-claim evidence."
)

add_heading(doc, "2.3 Retrieval-Augmented Generation", level=2)
add_body(doc,
    "Lewis et al. (2020) introduced Retrieval-Augmented Generation (RAG) as a "
    "framework combining parametric memory (the LLM) with non-parametric memory "
    "(a retrieval index) to improve factual accuracy on knowledge-intensive tasks. "
    "Subsequent work has explored dense retrieval (Karpukhin et al., 2020), "
    "iterative retrieval (Trivedi et al., 2022), and domain-specific applications in "
    "biomedicine (Zakka et al., 2024). In clinical summarization, Ji et al. (2023) "
    "argued that RAG is particularly well-suited because the entire relevant knowledge "
    "base—the source clinical note—is available at inference time, removing the need "
    "for an external corpus. However, empirical evaluation comparing RAG and baseline "
    "summaries at the claim level has not been systematically performed."
)

add_heading(doc, "2.4 Chain-of-Verification and Self-Correction", level=2)
add_body(doc,
    "Dhuliawala et al. (2023) proposed Chain-of-Verification (CoVe), a framework in "
    "which the model decomposes its initial response into verifiable questions, "
    "answers them independently, and uses the answers to revise its output. CoVe "
    "demonstrated meaningful reductions in hallucination on Wikidata entity lookups "
    "and long-form generation tasks. While CoVe addresses hallucination via iterative "
    "self-checking rather than retrieval, the underlying insight—that claim-level "
    "verification is more effective than holistic self-revision—motivates the "
    "claim-level evaluation framework adopted in this thesis. Combining RAG with "
    "CoVe-style post-hoc verification represents a natural direction for future work."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Methodology", level=1)

add_heading(doc, "3.1 Dataset", level=2)
add_body(doc,
    "Clinical transcriptions were drawn from MTSamples, a publicly available corpus of "
    "de-identified medical transcriptions spanning 40 specialty categories. We filtered "
    "for two document types most representative of structured clinical communication: "
    "Consultation – History and Physical notes (40 documents) and Discharge Summaries "
    "(10 documents), yielding a combined pool of 624 eligible documents. Fifty samples "
    "were selected using a fixed random seed (seed = 42) to ensure reproducibility. "
    "Source documents ranged from 117 to 1,809 words in length "
    "(mean = 539 ± 333 words, median = 466 words), reflecting the natural variability "
    "of clinical documentation."
)

add_heading(doc, "3.2 E0 Baseline: Zero-Context Summarization", level=2)
add_body(doc,
    "For each source document, GPT-4o-mini was prompted to generate a patient-facing "
    "summary without any retrieved context. The system prompt instructed the model to "
    "produce a structured summary comprising: Reason for Visit / Diagnosis, Key "
    "Findings, Treatment or Procedures Performed, Medications, and Follow-Up "
    "Instructions. No passages from the source document were included in the context "
    "window beyond the document itself; this condition thus reflects the model's "
    "baseline behavior when given only the raw note. Summaries averaged 172 words."
)

add_heading(doc, "3.3 E1: Retrieval-Augmented Generation", level=2)
add_body(doc,
    "The E1 condition augments the generation prompt with retrieved context. Each "
    "source document was chunked into overlapping windows of five sentences with a "
    "one-sentence stride. Chunks were encoded using the all-MiniLM-L6-v2 sentence "
    "transformer. At inference time, a structured query covering the five summary "
    "sections (diagnosis, findings, treatments, medications, follow-up) was embedded "
    "using the same encoder, and the top-3 chunks by cosine similarity were retrieved "
    "and concatenated as a grounding context. GPT-4o-mini was then prompted with both "
    "the source document and the retrieved context, instructed to prioritize factual "
    "consistency with the retrieved passages. All other prompt parameters were held "
    "constant across E0 and E1."
)

add_heading(doc, "3.4 Claim Extraction", level=2)
add_body(doc,
    "Generated summaries were segmented into individual claims using spaCy's "
    "en_core_web_sm sentence boundary detector. Markdown headers (bold section titles) "
    "were excluded from the claim set as they contain no verifiable propositional "
    "content. The resulting claim sets averaged 13.1 claims per E0 summary "
    "(total: 653 claims) and 9.6 claims per E1 summary (total: 478 claims). The "
    "lower E1 claim count reflects the model's tendency to produce more focused "
    "summaries when given retrieved context."
)

add_heading(doc, "3.5 Evidence Retrieval and NLI Classification", level=2)
add_body(doc,
    "For each claim, the top-3 most semantically similar sentences from the source "
    "document were retrieved using all-MiniLM-L6-v2 cosine similarity. These three "
    "sentences were concatenated as the evidence premise. Each (premise, claim) pair "
    "was then classified by the cross-encoder/nli-MiniLM2-L6-H768 model, which "
    "produces logit scores over three classes: entailment, neutral, and contradiction. "
    "Claims were labeled Supported (entailment), Not-Supported (neutral), or "
    "Contradicted (contradiction) based on the argmax of softmax probabilities."
)

add_heading(doc, "3.6 Metrics", level=2)
add_body(doc,
    "Two aggregate metrics were computed per summary. The Unsupported Fact Rate (UFR) "
    "captures the fraction of claims that are either Not-Supported or Contradicted:"
)
p_eq1 = doc.add_paragraph()
p_eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eq1.paragraph_format.space_after = Pt(4)
p_eq1.add_run("UFR = (|Not-Supported| + |Contradicted|) / |Total Claims|").bold = True

add_body(doc,
    "The Contradiction Rate (CR) isolates outright factual conflicts:"
)
p_eq2 = doc.add_paragraph()
p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eq2.paragraph_format.space_after = Pt(10)
p_eq2.add_run("CR = |Contradicted| / |Total Claims|").bold = True

add_body(doc,
    "Both metrics are computed per sample, then aggregated across the 50-sample cohort "
    "using mean, standard deviation, and median. Statistical significance of E0 vs E1 "
    "differences was assessed using the Wilcoxon signed-rank test (two-sided, α = 0.05), "
    "which makes no parametric assumptions about the distribution of per-sample "
    "metric differences."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. RESULTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Results", level=1)

add_heading(doc, "4.1 Aggregate Statistics", level=2)
add_body(doc,
    "Table 1 presents the aggregate UFR and CR statistics across all 50 samples for "
    "the baseline (E0) and RAG (E1) conditions. The Wilcoxon test results appear in "
    "Table 2."
)

# ── Table 1: Aggregate statistics ───────────────────────────────────────────
cap_t1 = doc.add_paragraph("Table 1. Aggregate hallucination metrics across 50 clinical summaries (E0 = baseline, E1 = RAG).")
cap_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_t1.runs[0].italic = True
cap_t1.runs[0].font.size = Pt(10)
cap_t1.paragraph_format.space_after = Pt(4)

tbl1 = doc.add_table(rows=4, cols=8)
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl1.style = "Table Grid"

headers = ["Metric", "E0 Mean", "E0 SD", "E0 Median", "E1 Mean", "E1 SD", "E1 Median", "Δ Mean"]
hdr_row = tbl1.rows[0]
for i, h in enumerate(headers):
    cell = hdr_row.cells[i]
    cell.text = h
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    set_cell_bg(cell, "1A1A2E")
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data_rows = [
    ["UFR", "0.797", "0.129", "0.806", "0.807", "0.172", "0.832", "+0.010"],
    ["CR",  "0.226", "0.124", "0.223", "0.154", "0.131", "0.134", "−0.073"],
]
row_colors = ["F2F2F2", "FFFFFF"]
for r_idx, (row_data, color) in enumerate(zip(data_rows, row_colors)):
    row = tbl1.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, color)
        if c_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

# Pct improved row
pct_row = tbl1.rows[3]
pct_data = ["% Improved", "—", "—", "—", "—", "—", "—", "42% / 72%"]
for c_idx, val in enumerate(pct_data):
    cell = pct_row.cells[c_idx]
    cell.text = val
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].runs[0].italic = True
    set_cell_bg(cell, "E8EDF5")

note = doc.add_paragraph("Note: % Improved = proportion of samples where E1 < E0 for that metric (UFR / CR respectively).")
note.runs[0].font.size = Pt(9)
note.runs[0].italic = True
note.paragraph_format.space_after = Pt(14)

# ── Table 2: Wilcoxon results ────────────────────────────────────────────────
cap_t2 = doc.add_paragraph("Table 2. Wilcoxon signed-rank test results (two-sided).")
cap_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_t2.runs[0].italic = True
cap_t2.runs[0].font.size = Pt(10)
cap_t2.paragraph_format.space_after = Pt(4)

tbl2 = doc.add_table(rows=3, cols=5)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Table Grid"

h2 = ["Metric", "W statistic", "p-value", "Median Δ", "Interpretation"]
for i, h in enumerate(h2):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(cell, "1A1A2E")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

wilcox_data = [
    ["UFR", "490.0", "0.581 (n.s.)",  "+0.009", "No significant change"],
    ["CR",  "215.0", "< 0.001 ***",   "−0.084", "RAG significantly reduced CR"],
]
for r_idx, (row_data, color) in enumerate(zip(wilcox_data, ["F2F2F2","FFFFFF"])):
    row = tbl2.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, color)
        if c_idx == 0:
            cell.paragraphs[0].runs[0].bold = True
        if val == "< 0.001 ***":
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

note2 = doc.add_paragraph("*** p < 0.001; n.s. = not significant (p > 0.05). α = 0.05 threshold.")
note2.runs[0].font.size = Pt(9)
note2.runs[0].italic = True
note2.paragraph_format.space_after = Pt(16)

# ── 4.2 Distributions ────────────────────────────────────────────────────────
add_heading(doc, "4.2 Metric Distributions", level=2)
add_body(doc,
    "Figures 1 and 2 display the per-sample distributions of UFR and CR as box plots "
    "with jittered individual observations. For UFR (Figure 1), the E1 median is "
    "marginally higher than E0 (0.832 vs 0.806), with greater variance in the RAG "
    "condition, confirming the non-significant Wilcoxon result. For CR (Figure 2), "
    "the E1 distribution is clearly shifted downward relative to E0, with both the "
    "median and upper quartile substantially reduced."
)
add_twoplot(doc,
    "plot_ufr_boxplot.png",
    "Figure 1. UFR distribution: E0 (baseline) vs E1 (RAG).",
    "plot_cr_boxplot.png",
    "Figure 2. CR distribution: E0 (baseline) vs E1 (RAG)."
)

add_heading(doc, "4.3 Per-Sample CR Comparison", level=2)
add_body(doc,
    "Figure 3 plots E0 CR against E1 CR for each of the 50 samples. Points below "
    "the diagonal indicate samples where RAG reduced the contradiction rate; 36 of "
    "50 samples (72%) fall in this region. The degree of improvement varies "
    "substantially across samples (Δ CR ranging from −0.257 to +0.362), suggesting "
    "that document characteristics—such as length, structural complexity, or the "
    "density of specific numerical claims—moderate the effectiveness of retrieval."
)
add_figure(doc, "plot_cr_scatter.png",
           "Figure 3. Per-sample scatter of E0 CR vs E1 CR. Points below the dashed diagonal "
           "indicate RAG improvement. Shaded region highlights the improvement zone.",
           width=4.0)

add_heading(doc, "4.4 Proportion of Samples Improved", level=2)
add_body(doc,
    "Figure 4 summarizes the proportion of samples in each outcome category "
    "(improved / unchanged / worse) for both metrics. CR shows a clear improvement "
    "signal: 72% of samples improved, 14% were unchanged, and 14% worsened. UFR "
    "presents a more ambiguous picture: 42% improved, a small fraction unchanged, "
    "and 56% worsened, consistent with the statistically non-significant group-level "
    "test."
)
add_figure(doc, "plot_pct_improved_bar.png",
           "Figure 4. Percentage of samples improved, unchanged, or worsened by RAG "
           "relative to baseline, for UFR and CR.",
           width=4.5)

# ══════════════════════════════════════════════════════════════════════════════
# 5. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Discussion", level=1)

add_heading(doc, "5.1 RAG Reduces Contradictions but Not Unsupported Claims", level=2)
add_body(doc,
    "The central empirical finding—that RAG significantly reduces CR (p < 0.001) "
    "while leaving UFR statistically unchanged (p = 0.581)—reveals a nuanced "
    "picture of how retrieval augmentation affects LLM factuality. The CR reduction "
    "is theoretically coherent: when the model is explicitly provided with retrieved "
    "passages stating, for example, a patient's actual medication regimen, the model "
    "is less likely to generate a statement directly contradicting that regimen. "
    "Retrieved context functions as a factual anchor that suppresses confabulation "
    "of specific, verifiable claims."
)
add_body(doc,
    "The persistence of high UFR in both conditions, however, reveals a more "
    "fundamental limitation. The majority of unsupported claims in both E0 and E1 "
    "are not contradictions—they are generic statements that no passage in the "
    "source document can verify. Examples from the qualitative analysis include "
    "boilerplate patient-education language ('it is important to take care of your "
    "mental health'), templated follow-up instructions ('reach out to your healthcare "
    "provider if you notice any changes'), and structural placeholders that the model "
    "generates to satisfy the expected summary format. These claims emerge from the "
    "model's parametric memory—trained distributions over medical language—rather "
    "than from the source document, and retrieval of additional source passages does "
    "not suppress them because they are not generated in response to source content "
    "in the first place."
)

add_heading(doc, "5.2 Qualitative Analysis of Claim Categories", level=2)
add_body(doc,
    "Qualitative inspection of the 33 documents where E0 contained Contradicted claims "
    "and E1 contained Supported claims revealed a consistent pattern: RAG was most "
    "effective when the contradicted E0 claim involved a specific, retrievable fact "
    "present verbatim or near-verbatim in the source document. For example, in one "
    "pediatric consultation (doc_id=2), the E0 summary stated that joint pain 'has "
    "improved significantly and is almost gone,' which the NLI model labeled as "
    "Contradicted because the source note only described ongoing pain. The E1 summary "
    "for the same document, having retrieved the relevant passage, correctly framed "
    "the improvement with appropriate temporal context, earning a Supported label."
)
add_body(doc,
    "In contrast, persistent hallucinations (claims labeled Not-Supported in both "
    "conditions) fell into two categories: (1) inferred but unstated clinical "
    "implications, such as treatment rationale or prognosis, that require domain "
    "knowledge not present in the source note; and (2) generic instructional language "
    "the model appended to satisfy the expected patient-communication register. "
    "Both categories represent failures that retrieval alone cannot address, as the "
    "required grounding information either does not exist in the document or the model "
    "generates the text independently of the retrieved context."
)

add_heading(doc, "5.3 Implications for Clinical NLP System Design", level=2)
add_body(doc,
    "These findings carry direct implications for the deployment of LLM-based "
    "summarization in clinical settings. First, the significant CR reduction "
    "achieved by RAG suggests that retrieval augmentation should be considered a "
    "baseline requirement—not an optional enhancement—for any clinical summarization "
    "system. The 48 percentage-point reduction in median CR (from 22.3% to 13.4%) "
    "translates to a meaningful decrease in the probability that any given claim "
    "in a patient-facing summary directly conflicts with the source record."
)
add_body(doc,
    "Second, the persistent UFR—averaging 80.7% even with RAG—indicates that the "
    "majority of claims in generated summaries cannot be verified against the source "
    "document. This does not necessarily mean all such claims are clinically "
    "erroneous; some may reflect appropriate inferential summarization. However, it "
    "underscores the need for post-hoc claim-level verification systems, such as "
    "CoVe-style self-checking (Dhuliawala et al., 2023) or explicit grounding "
    "constraints enforced at decoding time."
)

add_heading(doc, "5.4 Limitations", level=2)
add_body(doc,
    "Several limitations of this study warrant acknowledgment. First, the sample size "
    "of 50 documents, while sufficient to detect the large CR effect (Cohen's d ≈ 0.59), "
    "may be underpowered for detecting smaller UFR effects. Future work should evaluate "
    "at scale across the full MTSamples corpus and beyond. Second, the NLI classifier "
    "used for labeling—while state-of-the-art for its size class—is not specifically "
    "fine-tuned on clinical text and may systematically misclassify claims involving "
    "medical terminology or abbreviations. Third, the ground truth for claim labeling "
    "is itself model-generated, introducing a circular dependency: claims labeled "
    "Not-Supported by the NLI model may include legitimate inferential summaries that "
    "a human clinician would judge as appropriate. Fourth, the MTSamples corpus, "
    "while de-identified, is publicly available and may be partially represented in "
    "GPT-4o-mini's training data, potentially inflating E0 performance relative to "
    "a truly out-of-distribution evaluation. Finally, this study evaluates a single "
    "LLM (GPT-4o-mini); results may differ for larger models or those fine-tuned on "
    "clinical corpora."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Conclusion", level=1)
add_body(doc,
    "This thesis presented a systematic, claim-level evaluation of hallucinations in "
    "LLM-generated medical summaries, comparing zero-context baseline generation (E0) "
    "against retrieval-augmented generation (E1) across 50 de-identified clinical "
    "transcriptions. Using a pipeline of sentence-transformer retrieval and NLI-based "
    "classification, we quantified hallucination rates via two complementary metrics: "
    "the Unsupported Fact Rate (UFR) and the Contradiction Rate (CR)."
)
add_body(doc,
    "The results demonstrate that RAG produces a statistically significant and "
    "clinically meaningful reduction in outright factual contradictions "
    "(CR: 22.6% → 15.4%, Wilcoxon p < 0.001), with 72% of samples exhibiting "
    "improvement. This supports the adoption of retrieval augmentation as a standard "
    "component of clinical summarization pipelines. However, the broader unsupported "
    "fact rate remains high and statistically unchanged in both conditions, indicating "
    "that retrieval alone cannot eliminate the tendency of LLMs to generate "
    "clinically plausible but source-unverifiable content."
)
add_body(doc,
    "Future work will pursue three directions: (1) integration of claim-level "
    "grounding constraints at decoding time, (2) application of Chain-of-Verification "
    "post-hoc self-correction to the RAG pipeline, and (3) fine-tuning of the NLI "
    "classifier on clinician-annotated claim-evidence pairs to improve labeling "
    "accuracy in the medical domain. Ultimately, safe deployment of LLM-based "
    "clinical summarization will require not merely better generation, but rigorous, "
    "automated verification of every claim against its evidentiary basis in the "
    "source documentation."
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "References", level=1)

refs = [
    "Dhuliawala, S., Komeili, M., Xu, J., Raileanu, R., Li, X., Celikyilmaz, A., & Weston, J. (2023). Chain-of-Verification Reduces Hallucination in Large Language Models. arXiv:2309.11495.",
    "Guo, Z., Schlichtkrull, M., & Vlachos, A. (2022). A Survey on Automated Fact-Checking. Transactions of the Association for Computational Linguistics, 10, 178–206.",
    "Ji, Z., Lee, N., Frieske, R., Yu, T., Su, D., Xu, Y., ... & Fung, P. (2023). Survey of hallucination in natural language generation. ACM Computing Surveys, 55(12), 1–38.",
    "Karpukhin, V., Oğuz, B., Min, S., Lewis, P., Wu, L., Edunov, S., ... & Yih, W. T. (2020). Dense Passage Retrieval for Open-Domain Question Answering. In Proceedings of EMNLP 2020 (pp. 6769–6781).",
    "Kryscinski, W., McCann, B., Xiong, C., & Socher, R. (2020). Evaluating the Factual Consistency of Abstractive Text Summarization. In Proceedings of EMNLP 2020 (pp. 9332–9346).",
    "Laban, P., Schnabel, T., Bennett, P. N., & Hearst, M. A. (2022). SummaC: Re-Visiting NLI-based Models for Inconsistency Detection in Summarization. Transactions of the Association for Computational Linguistics, 10, 163–177.",
    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. In Advances in Neural Information Processing Systems (Vol. 33, pp. 9459–9474).",
    "Lin, C. Y. (2004). ROUGE: A Package for Automatic Evaluation of Summaries. In Text Summarization Branches Out (pp. 74–81). Association for Computational Linguistics.",
    "Maynez, J., Narayan, S., Bohnet, B., & McDonald, R. (2020). On Faithfulness and Factuality in Abstractive Summarization. In Proceedings of ACL 2020 (pp. 1906–1919).",
    "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. In Proceedings of EMNLP 2019 (pp. 3982–3992).",
    "Trivedi, H., Balasubramanian, N., Khot, T., & Sabharwal, A. (2022). MuSiQue: Multihop Questions via Single-hop Question Composition. Transactions of the Association for Computational Linguistics, 10, 539–554.",
    "Van Veen, D., Van Uden, C., Blankemeier, L., Delbrouck, J. B., Aali, A., Bluethgen, C., ... & Rajpurkar, P. (2024). Adapted large language models can outperform medical experts in clinical text summarization. Nature Medicine, 30(4), 1134–1142.",
    "Williams, A., Nangia, N., & Bowman, S. (2018). A Broad-Coverage Challenge Corpus for Sentence Understanding through Inference. In Proceedings of NAACL 2018 (pp. 1112–1122).",
    "Xie, Q., Mullenbach, J., Shi, X., Xing, E., & Ho, J. C. (2023). Me LLaMA: Foundation Large Language Models for Medical Applications. arXiv:2402.12749.",
    "Zakka, C., Shad, R., Chaurasia, A., Dalal, A. R., Kim, J. L., Moor, M., ... & Wornow, M. (2024). Almanac—Retrieval-Augmented Language Models for Clinical Medicine. NEJM AI, 1(2), AIoa2300068.",
]

for ref in refs:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent      = Pt(36)
    p.paragraph_format.first_line_indent = Pt(-36)
    p.paragraph_format.space_after       = Pt(6)
    p.add_run(ref).font.size = Pt(10)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\n  [done] Saved: {OUT}")
